from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\Reza\Desktop\promoguard-ai")
OUT = ROOT / "learning" / "project-mastery-guide" / "PromoGuard_Project_Mastery_Guide_FA.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
RED = "9B1C1C"
GOLD = "7A5A00"
GRAY = "5B6573"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


def set_font(run, name="Tahoma", size=11, color="1F2937", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_para(doc, text="", style=None, bold_prefix=None, color=None, size=11, space_after=6):
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=size, color=color or NAVY, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, size=size, color=color or "1F2937")
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=color or "1F2937")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_rtl(p)
    r = p.add_run(text)
    set_font(r, size={1: 18, 2: 14, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: NAVY}[level], bold=True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_rtl(p)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_callout(doc, label, text, fill=PALE, label_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_rtl(p)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "\n")
    set_font(r, size=11, color=label_color, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    widths = widths or [9360 // len(headers)] * len(headers)
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        set_rtl(p)
        r = p.add_run(head)
        set_font(r, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_rtl(p)
            r = p.add_run(str(value))
            set_font(r, size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def code(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(text.splitlines()):
        r = p.add_run(line)
        set_font(r, name="Consolas", size=9, color=NAVY)
        if idx < len(text.splitlines()) - 1:
            r.add_break()
    return p


def page_break(doc):
    doc.add_page_break()


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Tahoma"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Tahoma")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Tahoma")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = "Tahoma"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Tahoma")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Tahoma")
        style._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        style.font.color.rgb = RGBColor.from_string({1: BLUE, 2: BLUE, 3: NAVY}[level])
        style.font.size = Pt({1: 18, 2: 14, 3: 12}[level])
    header = sec.header.paragraphs[0]
    set_rtl(header)
    rh = header.add_run("PromoGuard  |  راهنمای تسلط و ارائه")
    set_font(rh, size=8.5, color=GRAY)
    footer = sec.footer.paragraphs[0]
    set_rtl(footer)
    rf = footer.add_run("راهنمای آموزشی داخلی — نسخهٔ پروژه: v0.5.2-park-demo")
    set_font(rf, size=8.5, color=GRAY)


def build():
    doc = Document()
    setup(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    r = p.add_run("PromoGuard")
    set_font(r, size=31, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("راهنمای کامل فهم، ارائه و دفاع از پروژه")
    set_font(r, size=19, color=BLUE, bold=True)
    add_para(doc, "از توضیح کودکانه تا معماری و آمار سطح فعلی پروژه", size=13, color=GRAY, space_after=20).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(doc, "هدف این جزوه", "بعد از خواندن این سند باید بتوانی در یک ارائهٔ یک‌ساعته توضیح بدهی PromoGuard چه مشکلی را حل می‌کند، داده از کجا می‌آید، هر ماژول چه می‌کند، چرا خروجی‌ها محافظه‌کارانه‌اند، چه چیزهایی هنوز اثبات نشده و به سؤال‌های فنی و کسب‌وکاری چگونه پاسخ بدهی.", fill=LIGHT)
    add_para(doc, "نسخهٔ فنی: v0.5.2-park-demo | دادهٔ کسب‌وکاری: واقعی و عمومی از dunnhumby | وضعیت: MVP قابل اجرا، شواهد عمومی آماده، ارسال رسمی نیازمند اطلاعات خصوصی و ویدیو", size=10, color=GRAY, space_after=8)
    page_break(doc)

    add_heading(doc, "فهرست مسیر مطالعه", 1)
    add_para(doc, "این سند عمداً پلکانی است. اگر زمان کم داشتی، بخش‌های ستاره‌دار را بخوان؛ اگر می‌خواهی واقعاً مسلط شوی، از ابتدا تا انتها جلو برو.")
    add_table(doc, ["مرحله", "موضوع", "نتیجه‌ای که باید بتوانی بگویی"], [
        ("۱", "داستان خیلی ساده", "PromoGuard مثل یک داور محتاط برای تخفیف‌هاست."),
        ("۲", "مسئله و ارزش", "چرا افزایش فروش الزاماً به معنی موفقیت پروموشن نیست."),
        ("۳", "داده و جریان کار", "CSV چگونه به پنل هفتگی و گزارش تبدیل می‌شود."),
        ("۴", "معماری و کد", "هر پوشه و ماژول چه مسئولیتی دارد."),
        ("۵", "آمار و عدم‌قطعیت", "baseline، WAPE، interval و محدودیت علیت."),
        ("۶", "محصول و داشبورد", "چطور خروجی typed به تجربهٔ کاربر تبدیل می‌شود."),
        ("۷", "تست و شواهد", "چرا ۷۳ تست و CI مهم‌اند، اما کافی برای اثبات بازار نیستند."),
        ("۸", "دفاع و پرسش‌ها", "جواب آماده برای سؤال‌های سخت داور یا استخدام‌کننده."),
    ], [700, 2600, 6060])
    add_callout(doc, "قانون طلایی ارائه", "هر بار که عدد یا ادعایی می‌گویی، بلافاصله بگو این ادعا دقیقاً چه چیزی را ثابت می‌کند و چه چیزی را ثابت نمی‌کند. این عادت، تفاوت یک پروژهٔ جدی آماری با یک داشبورد نمایشی است.", fill="FFF7E6", label_color=GOLD)

    add_heading(doc, "۱. توضیح پروژه برای یک نوجوان ۱۵ ساله", 1)
    add_heading(doc, "داستان یک تخفیف", 2)
    add_para(doc, "فرض کن یک فروشگاه برای یک بستهٔ غلات تخفیف می‌گذارد. در هفتهٔ تخفیف، ۱۸۵ بسته فروخته می‌شود. مدیر می‌گوید: «عالی! تخفیف باعث شد فروش بالا برود.» اما شاید اگر تخفیف وجود نداشت هم ۱۷۰ بسته فروخته می‌شد. شاید مردم خرید هفتهٔ بعد را جلو انداخته‌اند. شاید کالای رقیب تمام شده بود. شاید فروشگاه تبلیغ دیگری هم داشته است.")
    add_para(doc, "PromoGuard می‌گوید: اول آرام باشیم و سؤال درست بپرسیم. فروش واقعی را با یک حدس منظم از فروش عادی مقایسه می‌کنیم، میزان نامعلومی را نشان می‌دهیم و اگر داده کافی نیست، به‌جای ادعای قطعی می‌گوییم «شواهد بیشتری لازم است». ")
    add_table(doc, ["چیز ساده", "معنای فنی"], [
        ("فروش این هفته", "Observed units؛ چیزی که واقعاً در داده دیده‌ایم."),
        ("فروش عادی مورد انتظار", "Baseline؛ تخمین counterfactual ساده از گذشتهٔ بدون پروموشن."),
        ("فاصلهٔ بین آن دو", "Difference؛ نشانهٔ اولیه، نه اثر علّی قطعی."),
        ("بازهٔ شک", "Uncertainty interval؛ محدوده‌ای که نشان می‌دهد تخمین دقیق نیست."),
        ("چراغ هشدار", "Warning؛ مشکل داده یا تفسیر مثل forward-buy risk."),
        ("تصمیم", "Recommendation؛ آزمون کنترل‌شده، شواهد بیشتر یا بررسی/کم‌اولویت‌سازی."),
    ], [3800, 5560])
    add_callout(doc, "یک‌جمله‌ای", "PromoGuard نمی‌گوید «این تخفیف حتماً موفق بود»؛ می‌گوید «با داده‌ای که داریم، چقدر می‌توانیم به این نتیجه اعتماد کنیم و قدم بعدی چیست؟»", fill=LIGHT)

    add_heading(doc, "۲. متن آماده برای شروع ارائه", 1)
    add_para(doc, "«PromoGuard یک سامانهٔ evidence-first برای ممیزی پروموشن خرده‌فروشی است. مشکل این است که افزایش فروش در هفتهٔ تخفیف، به‌تنهایی ثابت نمی‌کند تخفیف ارزش اقتصادی ایجاد کرده؛ ممکن است تقاضا از آینده جلو کشیده شده، کالای دیگری جایگزین شده یا داده ناقص باشد. PromoGuard ابتدا کیفیت داده را کنترل می‌کند، سپس با تاریخچهٔ قبل از رویداد baseline می‌سازد، بازهٔ عدم‌قطعیت و هشدارها را نشان می‌دهد و اگر شواهد کافی نباشد، از توصیهٔ قطعی خودداری می‌کند و طراحی آزمون کنترل‌شده را پیشنهاد می‌دهد.»")
    add_para(doc, "بعد از این جمله، سریع یک رویداد واقعی را در Demo Mode نشان بده: فروش مشاهده‌شده ۱۲۸، baseline برابر ۱۸۵، اختلاف -۵۷، بازهٔ اختلاف از -۱۲۷ تا ۱۳ و تصمیم `needs_more_evidence` به‌دلیل هشدار `FORWARD_BUY_RISK`. این مثال نشان می‌دهد سیستم برای خوشحال‌کردن مدیر، نتیجهٔ مثبت جعل نمی‌کند.", bold_prefix="بعد از این جمله،")

    page_break(doc)
    add_heading(doc, "۳. مسئلهٔ کسب‌وکاری و تفاوت با داشبورد", 1)
    add_heading(doc, "مسئله دقیق چیست؟", 2)
    add_para(doc, "در خرده‌فروشی، شرکت‌ها برای تخفیف و کمپین پول خرج می‌کنند. گزارش ساده معمولاً می‌گوید فروش قبل، حین و بعد چقدر بوده است. اما تصمیم مهم‌تر این است: آیا فروش اضافه واقعاً از پروموشن آمده یا فقط جابه‌جایی زمانی/محصولی رخ داده است؟")
    add_bullet(doc, "افزایش همان هفته ممکن است خرید آینده را جلو بکشد (forward-buy).")
    add_bullet(doc, "فروش یک کالا ممکن است به قیمت از دست‌رفتن فروش کالای رقیب باشد (cannibalization).")
    add_bullet(doc, "نبود موجودی باعث می‌شود تقاضای واقعی را نبینیم.")
    add_bullet(doc, "تخفیف، تبلیغ، فصل، آب‌وهوا یا رویداد هم‌زمان می‌توانند عامل تغییر باشند.")
    add_bullet(doc, "اگر دادهٔ margin و هزینه نداریم، فروش بیشتر را نمی‌توان سود بیشتر نامید.")
    add_heading(doc, "PromoGuard چه چیزی اضافه می‌کند؟", 2)
    add_table(doc, ["داشبورد عادی", "PromoGuard"], [
        ("نمایش KPI", "نمایش KPI همراه با قرارداد داده و وضعیت کیفیت"),
        ("ممکن است نتیجهٔ مثبت را برجسته کند", "ممیزی محافظه‌کارانه و امکان abstain"),
        ("اغلب uncertainty را پنهان می‌کند", "interval و دامنهٔ ادعا را اجباری می‌کند"),
        ("منطق ممکن است در فایل گزارش گم شود", "policy نسخه‌دار و audit قابل بازتولید"),
        ("اثبات موفقیت را القا می‌کند", "برای اثر علّی، آزمایش کنترل‌شده را لازم می‌داند"),
    ], [4680, 4680])
    add_callout(doc, "جملهٔ دفاعی", "PromoGuard جای تصمیم‌گیرندهٔ انسانی نیست و وعدهٔ سود نمی‌دهد؛ یک لایهٔ غربالگری است که کمک می‌کند قبل از تصمیم پرهزینه بفهمیم داده و شواهد چقدر قابل اتکا هستند.", fill="FFF7E6", label_color=GOLD)

    add_heading(doc, "۴. دادهٔ واقعی و قرارداد داده", 1)
    add_heading(doc, "داده از کجا می‌آید؟", 2)
    add_para(doc, "پروژه روی دیتاست عمومی واقعی `Breakfast at the Frat` از dunnhumby اجرا شده است. این داده، دادهٔ تولیدشده توسط ما نیست و برای نمایش اثر تجاری ایران استفاده نمی‌شود؛ فقط نشان می‌دهد pipeline فنی روی یک پنل واقعی کار می‌کند.")
    add_table(doc, ["واقعیت freeze‌شده", "مقدار", "معنا"], [
        ("ردیف‌های داده", "۵۲۴٬۹۵۰", "تعداد ردیف‌های پنل پردازش‌شده"),
        ("سری‌ها", "۳٬۹۰۹", "ترکیب‌های store–UPC"),
        ("ردیف پروموشن", "۱۴۹٬۳۸۶", "ردیف‌هایی که در پنل پروموشن دارند"),
        ("ردیف grain تکراری", "۰", "هیچ store–UPC–date تکراری در قرارداد نهایی"),
        ("دادهٔ کسب‌وکاری مصنوعی", "خیر", "دادهٔ واقعی عمومی، نه شبیه‌سازی تجاری"),
    ], [3000, 1800, 4560])
    add_heading(doc, "قرارداد داده یعنی چه؟", 2)
    add_para(doc, "قرارداد داده مثل قوانین ورود به سیستم است. قبل از تحلیل، برنامه بررسی می‌کند شناسه‌های فروشگاه و کالا وجود دارند، تاریخ قابل خواندن است، واحد فروش معتبر است، grain یکتا است، پرچم پروموشن قابل تفسیر است و فایل از سقف حجم/ردیف عبور نکرده است. اگر شرط‌های پایه خراب باشند، تحلیل متوقف می‌شود؛ خروجی زیبا از دادهٔ خراب خطرناک است.")
    add_para(doc, "داده‌های raw و processed در Git commit نشده‌اند. مخزن فقط کد، قرارداد، گزارش فشرده و artifactهای قابل بازبینی را نگه می‌دارد. این هم از نظر حجم و هم حریم داده تصمیم درستی است.")
    add_callout(doc, "تفاوت fixture و evidence", "ممکن است تست‌های واحد از دادهٔ کوچک ساختگی استفاده کنند تا یک edge case را بررسی کنند؛ این دادهٔ آزمایشی است. اما عددهای ارائه و Demo از دادهٔ واقعی عمومی می‌آیند و هرگز fixture به‌عنوان شواهد کسب‌وکاری معرفی نمی‌شود.", fill=PALE)

    page_break(doc)
    add_heading(doc, "۵. نقشهٔ معماری از بالا به پایین", 1)
    add_para(doc, "کل سیستم را مثل یک کارخانه ببین:")
    add_table(doc, ["مرحلهٔ کارخانه", "ماژول/مسیر پروژه", "وظیفه"], [
        ("دریافت مواد", "`src/promoguard/data/`", "خواندن فایل و کنترل قرارداد"),
        ("تمیزکاری و شکل‌دهی", "ingestion / panel", "ساخت پنل هفتگی استاندارد"),
        ("پیش‌بینی مرجع", "`src/promoguard/forecasting/`", "ساخت baseline و ارزیابی زمان‌مند"),
        ("ممیزی رویداد", "`src/promoguard/insights/`", "مقایسهٔ observed و baseline، interval و warning"),
        ("بسته‌بندی سرویس", "`apps/api/`", "ارائهٔ قرارداد HTTP و JSON تایپ‌شده"),
        ("نمایش برای انسان", "`apps/dashboard/`", "Demo Mode فارسی و نمودار"),
        ("کنترل کیفیت", "`tests/`, `.github/workflows/`", "تست، lint و CI"),
        ("حافظه و آموزش", "`docs/`, `learning/`, `reports/`", "تصمیم‌ها، شواهد و راهنمای یادگیری"),
    ], [2600, 3000, 3760])
    add_heading(doc, "جریان واقعی اجرا", 2)
    code(doc, "CSV / workbook عمومی\n        ↓\nloader + data contract\n        ↓\ncanonical weekly panel\n        ↓\nquality checks ──(رد)──> خطای قابل فهم\n        ↓\ntime-aware baseline\n        ↓\npromotion episode detection\n        ↓\nPromotionAuditResult (منبع حقیقت اعداد)\n        ↓\nFastAPI JSON  +  Streamlit Demo Mode")
    add_callout(doc, "اصل معماری", "محاسبهٔ عددی فقط در هستهٔ دامنه انجام می‌شود. API و داشبورد نباید دوباره فرمول فروش، baseline یا اختلاف را حساب کنند؛ آن‌ها فقط نتیجهٔ typed را منتقل و نمایش می‌دهند. این کار جلوی اختلاف عدد بین backend و UI را می‌گیرد.", fill=LIGHT)

    add_heading(doc, "۶. توضیح پوشه‌ها و فایل‌ها", 1)
    add_table(doc, ["مسیر", "نقش", "سؤال رایج"], [
        ("`src/promoguard/data`", "loader، پنل و validation", "آیا دادهٔ ورودی قرارداد را رعایت می‌کند؟"),
        ("`src/promoguard/forecasting`", "مدل baseline و ارزیابی", "آیا آینده به گذشته نشت کرده؟"),
        ("`src/promoguard/insights`", "ممیزی و policy تصمیم", "آیا نتیجه observational است یا causal؟"),
        ("`apps/api/main.py`", "FastAPI entrypoint", "چه endpointهایی در اختیار مصرف‌کننده است؟"),
        ("`apps/api/contracts.py`", "Pydantic contracts", "JSON چه شکل و نوعی دارد؟"),
        ("`apps/dashboard/app.py`", "صفحهٔ اجرایی Streamlit", "کاربر چگونه Demo را می‌بیند؟"),
        ("`apps/dashboard/presentation.py`", "نگاشت متن و نمودار", "چطور متن فارسی از محاسبه جدا شده؟"),
        ("`demo/phase4_smoke.py`", "اجرای end-to-end قابل تکرار", "آیا API با دادهٔ واقعی smoke شده؟"),
        ("`tests/`", "تست‌های واحد/یکپارچه/آماری", "کدام قراردادها محافظت می‌شوند؟"),
        ("`reports/`", "خلاصهٔ machine-readable شواهد", "عددهای ارائه از کجا آمده‌اند؟"),
        ("`submission/park-application-1405/`", "بستهٔ ارائهٔ پارک", "چطور ادعا را به سند وصل کنیم؟"),
    ], [2700, 3000, 3660])

    add_heading(doc, "۷. منطق ورود داده و پنل هفتگی", 1)
    add_para(doc, "دادهٔ خام معمولاً به شکل تراکنش یا workbook است. سیستم آن را به یک جدول استاندارد تبدیل می‌کند که هر ردیف آن یک ترکیب روشن دارد: فروشگاه، کالا، تاریخ هفته، مقدار فروش و وضعیت پروموشن. این ترکیب را grain می‌نامیم.")
    add_para(doc, "اگر دو ردیف دقیقاً برای یک store، UPC و date وجود داشته باشد، معلوم نیست کدام را باید باور کنیم. پس duplicate grain یک خطای جدی است. همچنین تاریخ ناقص یا مقدار فروش نامعتبر قبل از مدل‌سازی باید متوقف شود.")
    add_number(doc, "فایل ورودی را از مسیر raw می‌خوانیم؛ raw را تغییر نمی‌دهیم.")
    add_number(doc, "ستون‌ها را به schema canonical نگاشت می‌کنیم.")
    add_number(doc, "تاریخ را parse و داده را بر اساس grain بررسی می‌کنیم.")
    add_number(doc, "پنل هفتگی را تولید می‌کنیم؛ هفته‌های بدون فروش را با احتیاط و طبق قرارداد مدیریت می‌کنیم.")
    add_number(doc, "گزارش کیفیت و provenance می‌نویسیم تا بعداً بدانیم چه چیزی، از کجا و با چه وضعیتی وارد شد.")

    page_break(doc)
    add_heading(doc, "۸. baseline و ارزیابی زمانی", 1)
    add_heading(doc, "baseline چیست؟", 2)
    add_para(doc, "baseline پاسخ ساده‌ای به این سؤال است: اگر پروموشن نبود، با توجه به گذشته چه میزان فروش انتظار داشتیم؟ PromoGuard در نسخهٔ فعلی ادعا نمی‌کند این پاسخ counterfactual واقعی و کامل است؛ baseline فقط مرجع شفاف برای غربالگری است.")
    add_para(doc, "برای هر سری store–UPC، تاریخچهٔ قبل از رویداد بررسی می‌شود. مهم‌ترین اصل این است که اطلاعات آینده وارد ساخت baseline نشود. اگر برای پیش‌بینی هفتهٔ قبل از رویداد از فروش بعد از رویداد استفاده کنیم، مدل تقلبی خوب به نظر می‌رسد؛ این همان leakage است.")
    add_heading(doc, "مدل‌های مقایسه‌شده", 2)
    add_table(doc, ["مدل", "توضیح ساده", "نتیجهٔ پروژه"], [
        ("Seasonal naive", "فروش مشابه در چرخهٔ ۵۲ هفته‌ای", "WAPE: 0.40046"),
        ("Recursive naive", "فروش هفتهٔ قبلی و تکرار قدم‌به‌قدم", "WAPE: 0.34828"),
        ("انتخاب مدل", "مدل بهتر باید با paired rows و protocol مشخص سنجیده شود", "نتیجهٔ منفی حفظ شده؛ پنهان نشده"),
    ], [2300, 3900, 3160])
    add_callout(doc, "نکتهٔ مهم", "در این داده، baseline فصلی از naive بازگشتی بهتر نشده است. این شکست، بخشی از شواهد علمی پروژه است. پروژه نتیجهٔ بد را حذف نمی‌کند تا داستان موفقیت بسازد؛ این دقیقاً رفتار قابل اعتماد است.", fill="FFF7E6", label_color=GOLD)
    add_heading(doc, "ارزیابی time-aware", 2)
    add_para(doc, "ارزیابی با ۶ پنجرهٔ expanding انجام شده است. تعداد ردیف‌های واجد شرایط غیرپروموشنی ۵۸٬۱۳۱ و تعداد ردیف‌های paired برابر ۴۱٬۵۱۶ است؛ coverage جفت‌شده ۷۱٫۴۲٪ است. paired یعنی فقط جایی مقایسه می‌کنیم که هر دو مدل پیش‌بینی معتبر دارند، تا مقایسه ناعادلانه نباشد.")
    add_para(doc, "WAPE را می‌توان این‌طور ساده فهمید: مجموع خطای قدرمطلق را بر مجموع مقدار واقعی تقسیم می‌کنیم. این معیار برای مقایسهٔ خطای کلی فروش مفید است، اما به‌تنهایی نمی‌گوید یک پروموشن سودآور بوده یا علت فروش چه بوده است.")

    add_heading(doc, "۹. ممیزی پروموشن و عدم‌قطعیت", 1)
    add_heading(doc, "PromotionAuditResult چیست؟", 2)
    add_para(doc, "این شیء typed خروجی مرکزی سیستم است. به‌جای اینکه داشبورد چند عدد پراکنده داشته باشد، یک نتیجهٔ ساختاریافته می‌گیرد: شناسهٔ رویداد، observed units، baseline point، baseline interval، difference point، difference interval، warnings، recommendation و claim boundary.")
    add_table(doc, ["فیلد", "در مثال Demo", "چطور توضیح بدهی"], [
        ("Observed", "۱۲۸", "فروشی که واقعاً در دوره دیده شده"),
        ("Baseline point", "۱۸۵", "مرجع نقطه‌ای از تاریخچهٔ قبل از رویداد"),
        ("Baseline interval", "۱۱۵ تا ۲۵۵", "بازهٔ عدم‌قطعیت baseline"),
        ("Difference point", "-۵۷", "فروش مشاهده‌شده ۵۷ واحد پایین‌تر از نقطهٔ baseline"),
        ("Difference interval", "-۱۲۷ تا ۱۳", "اثر مشاهده‌شده می‌تواند از افت تا افزایش کوچک باشد"),
        ("Recommendation", "needs_more_evidence", "هنوز برای تصمیم قاطع داده کافی نیست"),
        ("Warning", "FORWARD_BUY_RISK", "احتمال جلوکشیدن خرید از آینده"),
    ], [2600, 1900, 4860])
    add_heading(doc, "چرا interval مهم است؟", 2)
    add_para(doc, "اگر فقط بگوییم اختلاف -۵۷ است، کاربر ممکن است فکر کند عدد دقیق و قطعی است. اما interval از -۱۲۷ تا ۱۳ نشان می‌دهد داده هنوز اجازه نمی‌دهد با اطمینان بگوییم اثر منفی یا مثبت بوده است. پس recommendation منطقی، نیاز به شواهد بیشتر است.")
    add_heading(doc, "policy تصمیم", 2)
    add_para(doc, "policy `promoguard-observational-screening@1.0.0` سه مسیر اصلی دارد: `candidate_for_controlled_test` یعنی این رویداد ارزش طراحی آزمون دارد؛ `needs_more_evidence` یعنی قبل از تصمیم باید داده کامل‌تر شود؛ `deprioritize_and_investigate` یعنی فعلاً اولویت پایین‌تر یا بررسی علت لازم است. این policy نسخه‌دار است تا تغییر منطق در آینده قابل ردیابی باشد.")
    add_callout(doc, "فرق correlation و causation", "هم‌زمانی پروموشن و فروش بالا فقط correlation مشاهده‌شده است. برای causal effect باید گروه کنترل، randomization یا طراحی علّی معتبر و diagnostics داشته باشیم. نسخهٔ فعلی عمداً چنین ادعایی نمی‌کند.", fill="FDECEC", label_color=RED)

    page_break(doc)
    add_heading(doc, "۱۰. API و قرارداد نرم‌افزاری", 1)
    add_para(doc, "FastAPI لایهٔ سرویس است. وظیفه‌اش این نیست که منطق آماری را از نو بنویسد؛ وظیفه‌اش دریافت ورودی معتبر، صدا زدن هستهٔ دامنه، تبدیل نتیجه به JSON قراردادی و برگرداندن خطای قابل فهم است.")
    add_heading(doc, "چرا Pydantic مهم است؟", 2)
    add_para(doc, "Pydantic به خروجی شکل مشخص می‌دهد. مثلاً `observed_units` باید عدد باشد، `recommendation` باید یکی از گزینه‌های تعریف‌شده باشد و interval باید دو مرز داشته باشد. این کار جلوی JSON مبهم و ناسازگاری بین backend و frontend را می‌گیرد.")
    add_heading(doc, "مسیرهای مهم اجرا", 2)
    add_table(doc, ["مسیر", "کار", "نحوهٔ دفاع"], [
        ("`/health`", "بررسی زنده‌بودن سرویس و نسخه", "برای smoke و deployment check است، نه اعتبار آماری"),
        ("endpoint پنل/کیفیت", "دریافت یا بررسی پنل استاندارد", "داده قبل از audit باید قرارداد را پاس کند"),
        ("endpoint audit", "تولید PromotionAuditResult", "نتیجه observational و typed است"),
        ("OpenAPI docs", "نمایش قرارداد تعاملی", "مصرف‌کننده شکل ورودی/خروجی را می‌بیند"),
    ], [2200, 3300, 3860])
    add_heading(doc, "خطاها چگونه مدیریت می‌شوند؟", 2)
    add_bullet(doc, "فایل نامعتبر: پیام عمومی دربارهٔ ستون/دادهٔ لازم، بدون نمایش traceback.")
    add_bullet(doc, "دادهٔ بیش از سقف: توقف قبل از تحلیل برای حفاظت از منابع.")
    add_bullet(doc, "ورودی ناشناخته: رد با قرارداد تایپ‌شده.")
    add_bullet(doc, "مشکل داخلی: در محیط توسعه قابل ردیابی، در Demo بدون افشای مسیر شخصی.")
    add_callout(doc, "نکتهٔ استخدامی", "این بخش نشان می‌دهد پروژه فقط notebook نیست: قرارداد API، error boundary، health check و smoke test دارد.", fill=LIGHT)

    add_heading(doc, "۱۱. داشبورد Streamlit و Demo Mode", 1)
    add_para(doc, "داشبورد برای داور یا کاربر غیر فنی است. در حالت عادی می‌تواند مسیر داده یا کنترل‌های توسعه را داشته باشد؛ در `--demo` یک مسیر کوتاه و کنترل‌شده ارائه می‌کند تا داور با یک کلیک نتیجهٔ واقعی را ببیند.")
    add_heading(doc, "سه قدم Demo Mode", 2)
    add_number(doc, "بارگذاری و کنترل کیفیت: سیستم پنل واقعی را پیدا و قرارداد آن را بررسی می‌کند.")
    add_number(doc, "انتخاب رویداد نماینده: سیستم طبق قاعدهٔ deterministic رویداد را انتخاب می‌کند؛ کاربر بهترین نتیجه را دستی انتخاب نمی‌کند.")
    add_number(doc, "ممیزی و نمایش: observed، baseline، interval، warning و recommendation در کارت و نمودار نشان داده می‌شوند.")
    add_para(doc, "در Demo Mode مسیر محلی، traceback، نیاز به API key و محاسبهٔ دوبارهٔ عدد در UI وجود ندارد. اعداد از `PromotionAuditResult` می‌آیند و `presentation.py` فقط متن فارسی، رنگ و دادهٔ نمودار را نگاشت می‌کند.")
    add_heading(doc, "رفع خطای اخیر ModuleNotFoundError", 2)
    add_para(doc, "وقتی Streamlit با مسیر فایل اجرا می‌شود، گاهی ریشهٔ پروژه در import path نیست. بنابراین `app.py` ریشهٔ مخزن را با `Path(__file__).resolve().parents[2]` پیدا می‌کند و قبل از import کردن `apps.dashboard.presentation` به `sys.path` اضافه می‌کند. این راه‌حل به مسیر شخصی hard-code نشده است و اجرای cloneهای دیگر را هم حفظ می‌کند.")
    code(doc, ".\\.venv\\Scripts\\python.exe -m streamlit run apps/dashboard/app.py -- --demo")

    page_break(doc)
    add_heading(doc, "۱۲. تست‌ها و CI: چرا باید اعتماد کنیم؟", 1)
    add_para(doc, "تست یعنی قبل از اینکه انسان به خروجی اعتماد کند، ماشین چند بار قراردادها و edge caseها را امتحان کند. عدد ۷۳ تست به معنی «هیچ باگی وجود ندارد» نیست؛ یعنی ۷۳ بررسی تعریف‌شده پاس شده است.")
    add_table(doc, ["لایهٔ تست", "نمونهٔ چیزی که کنترل می‌کند"], [
        ("Unit", "توابع کوچک، نگاشت recommendation، warning و chart records"),
        ("Data contract", "ستون‌های لازم، تاریخ، grain، مقدار فروش و محدودیت‌ها"),
        ("Statistical", "تقسیم زمانی، leakage، paired evaluation و metric"),
        ("Integration", "اتصال API به هسته و قرارداد پاسخ"),
        ("Golden/output", "ثبات خروجی‌های نمونه و مرز ادعا"),
        ("Browser/manual", "اجرای واقعی Demo، RTL، نبود مسیر شخصی و نبود traceback"),
    ], [2400, 6960])
    add_heading(doc, "CI چه می‌کند؟", 2)
    add_para(doc, "GitHub Actions روی runner تمیز Ubuntu و Python 3.11 نصب می‌کند، dependencyها را می‌گیرد، Ruff را اجرا می‌کند و pytest را می‌زند. این مهم است چون فقط سیستم شخصی تو معیار نیست. آخرین CI مربوط به commit رفع import نیز سبز است.")
    add_heading(doc, "warning شناخته‌شده", 2)
    add_para(doc, "در تست FastAPI یک `StarletteDeprecationWarning` دربارهٔ مرز dependency با httpx دیده می‌شود. warning مانع پاس‌شدن تست نیست، اما پنهان نشده است. پاسخ حرفه‌ای این است: «آن را ثبت کرده‌ایم، اثرش را از خطای پروژه جدا کرده‌ایم و باید در maintenance dependency رفع شود.»")
    add_callout(doc, "سه سطح شواهد", "تست می‌گوید کد طبق قراردادهای نوشته‌شده رفتار می‌کند؛ اجرای دادهٔ واقعی می‌گوید مسیر روی یک dataset واقعی کار کرده؛ pilot مشتری می‌گوید ارزش کسب‌وکاری در محیط واقعی وجود دارد. پروژه دو مورد اول را دارد، مورد سوم هنوز انجام نشده است.", fill="FFF7E6", label_color=GOLD)

    add_heading(doc, "۱۳. وضعیت فعلی پروژه و releaseها", 1)
    add_table(doc, ["نسخه/گیت", "چه چیزی تمام شده", "چه چیزی هنوز نیست"], [
        ("v0.5.1-foundation-correctness", "coverage، policy نسخه‌دار، real-data evidence و API smoke", "اثبات اثر تجاری"),
        ("v0.5.2-park-demo", "Demo Mode، نمودار interval، UI فارسی و privacy review", "ویدیوی نهایی و private form"),
        ("commit f829efa", "رفع اجرای مستقیم Streamlit و مستندات آن", "deploy عمومی"),
        ("مرحلهٔ ۵.۴", "فعال: فرم خصوصی، PDFها و ویدیو", "نیازمند اقدام صاحب پروژه"),
    ], [2500, 4300, 2560])
    add_para(doc, "درخواست از پارک هنوز نباید بر اساس مشتری، درآمد، patent، award یا اثر ایران ادعا بسازد. ارزش فعلی، توان ساخت یک MVP قابل ممیزی و درخواست دسترسی به شریک pilot است.")

    page_break(doc)
    add_heading(doc, "۱۴. محدودیت‌ها را چگونه توضیح بدهی؟", 1)
    add_para(doc, "محدودیت ضعف پنهان نیست؛ بخشی از قرارداد علمی محصول است. پاسخ خوب، محدودیت را می‌گوید و بعد قدم بعدی برای کاهش آن را مشخص می‌کند.")
    add_table(doc, ["محدودیت فعلی", "چرا مهم است", "قدم بعدی"], [
        ("دادهٔ public و observational", "اثر علّی و customer impact را ثابت نمی‌کند", "pilot با گروه کنترل/طراحی آزمایش"),
        ("margin و inventory نداریم", "فروش بیشتر را نمی‌توان profit نامید", "اتصال دادهٔ هزینه، margin و موجودی با مجوز"),
        ("LLM در runtime نیست", "متن آزاد ممکن است hallucination بدهد", "بعداً LLM محدود به JSON typed + golden eval"),
        ("authentication/rate limit کامل نیست", "برای production multi-tenant کافی نیست", "زیرساخت production پس از اعتبارسنجی مسئله"),
        ("تیم یک‌نفره", "ریسک ظرفیت و domain knowledge وجود دارد", "بعد از discovery جذب domain expert یا engineer"),
        ("benchmark اقتصادی نداریم", "ROI و willingness-to-pay نامعلوم است", "مصاحبه و pilot واقعی"),
    ], [2500, 3500, 3360])
    add_heading(doc, "چرا LLM را فعلاً اضافه نکردیم؟", 2)
    add_para(doc, "چون LLM برای عددسازی منبع حقیقت خوبی نیست. معماری درست این است: محاسبهٔ عددی deterministic و typed بماند؛ LLM فقط همین JSON را به زبان ساده توضیح دهد؛ سپس با تست طلایی بررسی کنیم عددها، warningها و کلمات عدم‌قطعیت تغییر نکرده‌اند. اگر API قطع شد، MVP باید همچنان کار کند.")
    add_callout(doc, "جملهٔ قوی برای مصاحبه", "من AI را به‌خاطر مد روز به محصول نچسباندم. اول correctness و claim boundary را ساختم؛ بعد اگر narrative generation اضافه شود، مدل فقط سخنگوی یک نتیجهٔ تایپ‌شده است، نه تولیدکنندهٔ حقیقت.", fill=LIGHT)

    add_heading(doc, "۱۵. سؤال‌های احتمالی و جواب آماده", 1)
    questions = [
        ("این فقط یک داشبورد نیست؟", "داشبورد KPI نشان می‌دهد؛ PromoGuard قبل از نتیجه‌گیری قرارداد داده، baseline زمانی، interval، warning و مرز علیت را اجرا می‌کند و می‌تواند abstain کند."),
        ("آیا ثابت کرده‌ای تخفیف فروش را زیاد کرده؟", "خیر. دادهٔ فعلی observational است. اختلاف مشاهده‌شده screening است، نه causal effect. برای ادعای اثر، آزمون کنترل‌شده و diagnostics لازم است."),
        ("چرا عدد baseline را باور کنیم؟", "آن را حقیقت نهایی نمی‌دانیم؛ یک مرجع شفاف از تاریخچهٔ قبل از رویداد است. بازهٔ عدم‌قطعیت و محدودیت‌ها کنار آن نمایش داده می‌شوند و مدل با ارزیابی زمان‌مند سنجیده شده است."),
        ("چرا نتیجهٔ Demo منفی است؟", "چون رویداد به‌صورت deterministic و غیرچری‌پیک انتخاب شده است. سیستم برای ساختن داستان مثبت، بهترین نمونه را انتخاب نکرده؛ `needs_more_evidence` خروجی علمی صادقانه است."),
        ("چرا ۷۳ تست کافی است؟", "کافی برای نشان‌دادن discipline مهندسی و قراردادهای فعلی است، نه اثبات نبود همهٔ خطاها یا ارزش بازار. هر تست دامنهٔ مشخص دارد."),
        ("آیا دادهٔ شما synthetic است؟", "دادهٔ شواهد از پنل عمومی واقعی dunnhumby است. ممکن است fixture کوچک در تست باشد، اما آن را به‌عنوان evidence کسب‌وکاری ارائه نمی‌کنیم."),
        ("چرا profit نشان نمی‌دهی؟", "چون margin، هزینهٔ تخفیف و inventory در dataset عمومی نیست. فروش بیشتر را بدون این‌ها profit نامیدن از نظر علمی نادرست است."),
        ("مشتری داری؟", "خیر. MVP روی دادهٔ عمومی واقعی ساخته شده. درخواست من از پارک، معرفی شریک discovery/pilot است، نه ادعای مشتری یا traction."),
        ("چطور به شغل Senior Data Scientist مربوط است؟", "مسئله را از data contract تا model evaluation، API، dashboard، tests، CI و monitoring boundary end-to-end دیده‌ام؛ در کنار آن limitation و causal inference را هم صریح نگه داشته‌ام."),
        ("اگر حجم داده ۱۰۰ برابر شود چه؟", "ابتدا profiling و benchmark می‌کنم؛ سپس پنل/warehouse، partition، query engine و orchestration را بر اساس bottleneck واقعی اضافه می‌کنم، نه به‌عنوان تزئین معماری."),
        ("چرا از deep learning استفاده نکردی؟", "مدل ساده و قابل توضیح برای baseline فعلی مناسب‌تر است. پیچیدگی فقط وقتی پذیرفته می‌شود که validation زمان‌مند نشان دهد ارزش افزوده دارد."),
        ("اگر یک مدیر بگوید فقط نتیجهٔ مثبت را بده؟", "سیستم باید claim boundary و warning را حفظ کند. تصمیم تجاری انسان است، اما نرم‌افزار نباید دادهٔ ناکافی را به موفقیت قطعی تبدیل کند."),
    ]
    for q, a in questions:
        add_heading(doc, q, 2)
        add_para(doc, a)

    page_break(doc)
    add_heading(doc, "۱۶. سناریوی ارائهٔ یک‌ساعته", 1)
    add_para(doc, "لازم نیست یک ساعت بدون توقف کد بخوانی. ارائه باید از داستان به شواهد و بعد به عمق فنی حرکت کند.")
    add_table(doc, ["زمان", "چه بگویی/نشان بدهی", "هدف"], [
        ("۰–۵ دقیقه", "مسئلهٔ تخفیف و تفاوت فروش بیشتر با سود/اثر واقعی", "جذب مخاطب"),
        ("۵–۱۰ دقیقه", "جملهٔ جایگاه‌یابی و کاربران هدف", "تعریف دقیق محصول"),
        ("۱۰–۲۰ دقیقه", "Demo Mode روی دادهٔ واقعی؛ رویداد، کارت و نمودار", "شاهد قابل مشاهده"),
        ("۲۰–۳۰ دقیقه", "pipeline داده تا PromotionAuditResult", "فهم معماری"),
        ("۳۰–۴۰ دقیقه", "baseline، WAPE، paired coverage و interval", "عمق آماری"),
        ("۴۰–۴۸ دقیقه", "API، contracts، presentation separation و import fix", "بلوغ مهندسی"),
        ("۴۸–۵۵ دقیقه", "تست‌ها، CI، محدودیت‌ها و roadmap", "اعتماد و صداقت"),
        ("۵۵–۶۰ دقیقه", "سؤال‌ها و درخواست دقیق از پارک/کارفرما", "جمع‌بندی و next step"),
    ], [1400, 5260, 2800])
    add_heading(doc, "دموی پیشنهادی", 2)
    add_number(doc, "ابتدا بگو: «این پنل public و واقعی است؛ ادعای customer impact ندارم.»")
    add_number(doc, "دستور Demo Mode را اجرا کن و فقط دکمهٔ اصلی را بزن.")
    add_number(doc, "شناسهٔ رویداد و نتیجهٔ ۱۲۸ در برابر ۱۸۵ را نشان بده.")
    add_number(doc, "به interval از -۱۲۷ تا ۱۳ اشاره کن و بگو چرا تصمیم قطعی نمی‌دهیم.")
    add_number(doc, "warning forward-buy را توضیح بده.")
    add_number(doc, "از داشبورد به کد/گزارش برو و نشان بده عدد از typed result آمده است.")
    add_number(doc, "با محدودیت‌ها تمام کن: بدون pilot، سود و علیت ادعا نمی‌شود.")
    add_callout(doc, "اگر Demo خراب شد", "از آخرین screenshot و artifact واقعی استفاده کن، ولی خرابی را پنهان نکن. یک MVP حرفه‌ای مسیر بازتولید و محدودیت دارد؛ حفظ صداقت بهتر از improvisation عددی است.", fill="FDECEC", label_color=RED)

    add_heading(doc, "۱۷. فرمان‌های مهم برای تمرین", 1)
    code(doc, "Set-Location C:\\Users\\Reza\\Desktop\\promoguard-ai\n.\\.venv\\Scripts\\python.exe -m streamlit run apps/dashboard/app.py -- --demo\n\n.\\.venv\\Scripts\\python.exe -m pytest -q\n.\\.venv\\Scripts\\python.exe -m ruff check .\n.\\.venv\\Scripts\\python.exe -m compileall -q src apps demo")
    add_para(doc, "برای API نیز می‌توانی سرویس را جداگانه بالا بیاوری و `/docs` را باز کنی. در ارائه، مهم‌تر از حفظ کردن تمام endpointها این است که بدانی API لایهٔ ارائه است و منطق عددی در هستهٔ domain قرار دارد.")

    page_break(doc)
    add_heading(doc, "۱۸. فرهنگ واژه‌ها به زبان خیلی ساده", 1)
    add_table(doc, ["واژه", "معنای ساده", "پاسخ کوتاه در جلسه"], [
        ("SKU / UPC", "شناسهٔ کالا", "برای ساخت سری فروشگاه–کالا لازم است."),
        ("Grain", "دقیقاً یک ردیف نمایندهٔ چیست", "اینجا store–UPC–date است."),
        ("Panel", "جدول مرتب‌شدهٔ چند هفته/سری", "واحد پایهٔ تحلیل زمانی."),
        ("Promotion episode", "دورهٔ پیوستهٔ پروموشن", "چند هفتهٔ رویداد به‌صورت یک episode."),
        ("Baseline", "حدس فروش عادی", "مرجع شفاف، نه حقیقت علّی."),
        ("Leakage", "دیدن جواب آینده هنگام آموزش", "با split زمانی جلوگیری می‌کنیم."),
        ("WAPE", "درصد خطای وزنی فروش", "برای مقایسهٔ خطای پیش‌بینی."),
        ("Interval", "بازهٔ عدم‌قطعیت", "نشان می‌دهد تخمین چقدر دقیق نیست."),
        ("Observational", "فقط چیزی که در داده دیده‌ایم", "بدون randomization، causal نیست."),
        ("Causal inference", "بررسی اثر علت", "نیازمند طراحی و فرض‌های معتبر."),
        ("Abstention", "امتناع از نتیجهٔ قطعی", "وقتی شواهد کافی نیست، توقف می‌کنیم."),
        ("Typed result", "خروجی با شکل و نوع مشخص", "API و UI یک منبع عددی مشترک دارند."),
        ("RAG", "بازیابی متن/داده برای مدل", "در فاز فعلی runtime نیست."),
        ("Agent", "سیستم چندمرحله‌ای تصمیم‌یار", "بعداً و فقط با guardrail قابل افزودن است."),
        ("CI", "تست خودکار روی runner تمیز", "کد فقط روی لپ‌تاپ اعتبارسنجی نمی‌شود."),
    ], [2100, 3900, 3360])

    add_heading(doc, "۱۹. برنامهٔ ادامهٔ منطقی", 1)
    add_para(doc, "بعد از بستهٔ پارک، توسعه باید مرحله‌ای باشد. هر مرحله باید با evidence و gate تمام شود، نه با اضافه‌کردن ابزارهای مد روز.")
    add_table(doc, ["فاز آینده", "هدف", "شرط ورود/خروج"], [
        ("Phase 6", "benchmark روی آزمایش random واقعی و روش‌های causal", "تعریف treatment/outcome و diagnostics کامل"),
        ("Phase 7", "cannibalization و forward-buy عمیق‌تر", "نمونهٔ عمومی و error rate مستند"),
        ("Phase 8", "بهینه‌سازی سناریو با margin/inventory واقعی", "constraint و human approval"),
        ("Phase 9", "warehouse، orchestration، monitoring و LLM محدود", "MVP بدون LLM همچنان سالم بماند"),
        ("Phase 10", "case study و role mapping", "داور در ۶۰ ثانیه ارزش و شواهد را بفهمد"),
    ], [1900, 4160, 3300])
    add_callout(doc, "اولویت واقعی", "اگر زمان یا منابع کم شد، data validation، time-aware evaluation، uncertainty، reproducibility و limitation را حذف نکن. styling، مدل پیچیده، LLM و optimization بعد از این‌ها هستند.", fill=LIGHT)

    add_heading(doc, "۲۰. چک‌لیست تسلط شخصی", 1)
    for item in [
        "می‌توانم پروژه را در یک جمله بدون jargon توضیح دهم.",
        "می‌توانم بگویم چرا فروش بیشتر مساوی سود یا اثر علّی نیست.",
        "می‌دانم دادهٔ واقعی از کجا آمده و چه چیزهایی در Git نیست.",
        "می‌توانم مسیر CSV تا PromotionAuditResult را روی کاغذ بکشم.",
        "می‌دانم leakage چیست و چرا split زمانی لازم است.",
        "می‌توانم observed=128، baseline=185، difference=-57 و interval را توضیح دهم.",
        "می‌دانم چرا recommendation `needs_more_evidence` است.",
        "می‌دانم API و dashboard نباید محاسبهٔ عدد را تکرار کنند.",
        "می‌توانم معنی ۷۳ تست، CI سبز و warning dependency را جداگانه توضیح دهم.",
        "می‌توانم حداقل پنج محدودیت و قدم بعدی هرکدام را بگویم.",
        "می‌دانم چه ادعاهایی را نباید بکنم: مشتری، درآمد، patent، ROI و اثر ایران.",
        "می‌توانم Demo Mode را بدون لو دادن مسیر شخصی اجرا کنم.",
    ]:
        add_bullet(doc, "☐ " + item)

    add_heading(doc, "جمع‌بندی نهایی برای حفظ کردن", 1)
    add_callout(doc, "نسخهٔ ۳۰ ثانیه‌ای", "PromoGuard یک MVP تحلیل داده برای ممیزی پروموشن است. روی دادهٔ واقعی عمومی، کیفیت داده و پنل هفتگی را کنترل می‌کند، baseline زمان‌مند می‌سازد، observed را با baseline همراه با interval و warning مقایسه می‌کند و اگر شواهد کافی نباشد abstain می‌کند. نسخهٔ فعلی اثر علّی، سود یا مشتری ایرانی را ادعا نمی‌کند؛ ارزش آن در reproducibility، قراردادهای typed، تست، CI و تصمیم‌گیری evidence-first است.", fill=NAVY, label_color="FFFFFF")
    add_para(doc, "منبع‌های مطالعهٔ بیشتر: `README.md`، `ROADMAP.md`، گزارش‌های `reports/`، بستهٔ `submission/park-application-1405/` و راهنماهای فارسی داخل `learning/`. عددهای ارائه باید همیشه از گزارش‌های نسخه‌دار خوانده شوند، نه از حافظه.", size=10, color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "PromoGuard — راهنمای کامل فهم، ارائه و دفاع از پروژه"
    doc.core_properties.subject = "راهنمای فارسی آموزشی و ارائه‌ای پروژه PromoGuard"
    doc.core_properties.author = "PromoGuard Project"
    doc.core_properties.keywords = "PromoGuard, data science, retail, promotion audit, Persian"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
